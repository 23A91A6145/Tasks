"""Knowledge ingestion & retrieval (tenant-isolated RAG).

Pipeline:  file/url/faq → parse text → chunk → embed → vector store
Retrieval: query → embed → search → ranked chunks (namespace = workspace id)
"""

import os
import re
import tempfile
from typing import Optional

from fastapi import HTTPException
from sqlalchemy import select
from sqlalchemy.orm import Session

from ..core.config import settings
from ..models import KnowledgeDocument, KnowledgeTag, Organization, document_tags
from .embeddings import get_embedder
from .vector import VectorPoint, get_vector_store

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt", ".csv", ".json"}
MAX_CHARS = 2_000_000


# ─────────────────────────── parsing ───────────────────────────


def extract_text(filename: str, content: bytes) -> str:
    """Parse an uploaded file into plain text based on its extension."""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(__import__("io").BytesIO(content))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    if ext == ".docx":
        from docx import Document

        doc = Document(__import__("io").BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n\n".join(parts)
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("utf-8", errors="replace")


def extract_url(url: str) -> str:
    from ..core.urlsafety import validate_public_url
    from fastapi import HTTPException

    try:
        url = validate_public_url(url)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc

    import trafilatura

    fetched = trafilatura.fetch_url(url)
    if not fetched:
        raise HTTPException(status_code=422, detail="Could not fetch that URL")
    text = trafilatura.extract(fetched)
    if not text:
        raise HTTPException(status_code=422, detail="No readable text found at that URL")
    return text


# ─────────────────────────── chunking ───────────────────────────


def _persist_text(organization_id: str, doc_id: str, text: str) -> str:
    """Save the extracted text so jobs can re-index a document offline."""
    directory = os.path.join(settings.STORAGE_DIR, "documents", organization_id)
    os.makedirs(directory, exist_ok=True)
    path = os.path.join(directory, f"{doc_id}.txt")
    with open(path, "w", encoding="utf-8") as handle:
        handle.write(text)
    return path


def load_persisted_text(doc: KnowledgeDocument) -> str:
    """Return the original extracted text for a document."""
    if doc.stored_path and os.path.exists(doc.stored_path):
        with open(doc.stored_path, "r", encoding="utf-8") as handle:
            return handle.read()
    return ""


def chunk_text(text: str, size: int | None = None, overlap: int | None = None) -> list[str]:
    """Recursive character-level splitter with overlap."""
    size = size or settings.CHUNK_SIZE
    overlap = overlap or settings.CHUNK_OVERLAP
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) <= size:
        return [text] if text else []

    chunks: list[str] = []
    for paragraph in re.split(r"\n{2,}", text):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        while len(paragraph) > size:
            split_at = paragraph.rfind(" ", 0, size)
            split_at = split_at if split_at > size * 0.5 else size
            chunks.append(paragraph[:split_at].strip())
            paragraph = paragraph[split_at - overlap :]
        if paragraph:
            chunks.append(paragraph.strip())
    return [c for c in chunks if c]


# ─────────────────────────── ingestion ───────────────────────────


def ingest_text(
    db: Session,
    *,
    organization: Organization,
    filename: str,
    text: str,
    source_type: str = "upload",
    uploaded_by_id: Optional[str] = None,
    source_url: Optional[str] = None,
    file_type: Optional[str] = None,
    tags: list[str] | None = None,
    doc_id: Optional[str] = None,
) -> KnowledgeDocument:
    """Embed and store document chunks under the tenant's vector namespace."""
    text = text[:MAX_CHARS]
    doc = db.execute(
        select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id)
    ).scalar_one_or_none() if doc_id else None

    if doc is None:
        doc = KnowledgeDocument(
            organization_id=organization.id,
            uploaded_by_id=uploaded_by_id,
            filename=filename,
            source_type=source_type,
            file_type=file_type or _type_from_name(filename),
            source_url=source_url,
            size_bytes=len(text.encode("utf-8")),
            status="processing",
        )
        db.add(doc)
        db.flush()

    doc.status = "processing"
    doc.error = None

    doc.stored_path = _persist_text(organization.id, doc.id, text)

    namespace = organization.id
    chunks = chunk_text(text)
    embedder = get_embedder()
    store = get_vector_store()
    vectors = embedder.embed(chunks) if chunks else []

    store.delete_document(namespace, doc.id)
    points = [
        VectorPoint(
            id=f"{doc.id}:{i}",
            document_id=doc.id,
            chunk_index=i,
            text=chunk,
            vector=vector,
            metadata={"filename": filename, "source_type": source_type},
        )
        for i, (chunk, vector) in enumerate(zip(chunks, vectors))
    ]
    store.upsert(namespace, points)

    doc.chunk_count = len(points)
    doc.status = "ready" if points else "failed"
    if not points:
        doc.error = "No searchable text extracted from this document."

    if tags:
        doc.tags = _ensure_tags(db, organization.id, tags)

    db.flush()
    return doc


def ingest_file(
    db: Session,
    *,
    organization: Organization,
    filename: str,
    content: bytes,
    uploaded_by_id: Optional[str] = None,
    tags: list[str] | None = None,
) -> KnowledgeDocument:
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=422,
            detail=f"Unsupported file type '{ext or 'unknown'}'. Allowed: {sorted(ALLOWED_EXTENSIONS)}",
        )
    if len(content) > settings.MAX_UPLOAD_MB * 1024 * 1024:
        raise HTTPException(
            status_code=413,
            detail=f"File exceeds the {settings.MAX_UPLOAD_MB} MB upload limit",
        )
    text = extract_text(filename, content)
    return ingest_text(
        db,
        organization=organization,
        filename=filename,
        text=text,
        source_type="upload",
        uploaded_by_id=uploaded_by_id,
        file_type=ext.lstrip("."),
        tags=tags,
    )


def ingest_url(
    db: Session,
    *,
    organization: Organization,
    url: str,
    uploaded_by_id: Optional[str] = None,
    tags: list[str] | None = None,
) -> KnowledgeDocument:
    text = extract_url(url)
    filename = url.rstrip("/").split("/")[-1] or "webpage"
    return ingest_text(
        db,
        organization=organization,
        filename=f"{filename}.md",
        text=text,
        source_type="url",
        uploaded_by_id=uploaded_by_id,
        file_type="url",
        source_url=url,
        tags=tags,
    )


def ingest_faq(
    db: Session,
    *,
    organization: Organization,
    name: str,
    content: str,
    uploaded_by_id: Optional[str] = None,
    tags: list[str] | None = None,
) -> KnowledgeDocument:
    return ingest_text(
        db,
        organization=organization,
        filename=f"{name}.md",
        text=content,
        source_type="faq",
        uploaded_by_id=uploaded_by_id,
        file_type="md",
        tags=tags,
    )


def _type_from_name(filename: str) -> str:
    return os.path.splitext(filename)[1].lstrip(".") or "txt"


# ─────────────────────────── tags ───────────────────────────


def _ensure_tags(db: Session, organization_id: str, names: list[str]) -> list[KnowledgeTag]:
    tags: list[KnowledgeTag] = []
    for name in dict.fromkeys(n.strip() for n in names if n.strip()):
        tag = db.execute(
            select(KnowledgeTag).where(
                KnowledgeTag.organization_id == organization_id,
                KnowledgeTag.name == name,
            )
        ).scalar_one_or_none()
        if tag is None:
            tag = KnowledgeTag(organization_id=organization_id, name=name)
            db.add(tag)
            db.flush()
        tags.append(tag)
    return tags


def list_tags(db: Session, organization_id: str) -> list[KnowledgeTag]:
    return db.execute(
        select(KnowledgeTag)
        .where(KnowledgeTag.organization_id == organization_id)
        .order_by(KnowledgeTag.name)
    ).scalars().all()


# ─────────────────────────── retrieval ───────────────────────────


def search(
    db: Session,
    organization_id: str,
    query: str,
    top_k: int = 0,
    filename_filter: Optional[str] = None,
) -> list[dict]:
    top_k = top_k or settings.SEARCH_TOP_K
    embedder = get_embedder()
    store = get_vector_store()
    vector = embedder.embed([query])[0]
    hits = store.search(organization_id, vector, top_k)
    results = [
        {
            "id": hit.id,
            "document_id": hit.document_id,
            "chunk_index": hit.chunk_index,
            "text": hit.text,
            "score": hit.score,
            "filename": hit.metadata.get("filename", ""),
            "source_type": hit.metadata.get("source_type", ""),
        }
        for hit in hits
    ]
    if filename_filter:
        results = [r for r in results if filename_filter in r["filename"]]
    return results


def delete_document(db: Session, organization_id: str, doc_id: str) -> KnowledgeDocument:
    doc = db.execute(
        select(KnowledgeDocument).where(
            KnowledgeDocument.id == doc_id,
            KnowledgeDocument.organization_id == organization_id,
        )
    ).scalar_one_or_none()
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found in this workspace")
    get_vector_store().delete_document(organization_id, doc.id)
    if doc.stored_path and os.path.exists(doc.stored_path):
        try:
            os.remove(doc.stored_path)
        except OSError:  # pragma: no cover
            pass
    db.delete(doc)
    db.flush()
    return doc
